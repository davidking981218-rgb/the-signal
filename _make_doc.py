"""THE SIGNAL 설명서 docx 생성 — 전면 개정본 (임베딩 마이그레이션 반영)"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 기본 폰트
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)


def H1(text):
    doc.add_heading(text, level=1)


def H2(text):
    doc.add_heading(text, level=2)


def H3(text):
    doc.add_heading(text, level=3)


def P(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "맑은 고딕"


def BULLET(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "맑은 고딕"


def CODE(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x60)


# ── 표지 ────────────────────────────────────────
title = doc.add_heading("THE SIGNAL — AI Daily Briefing", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("프로그램 설명서 (최신 개정본 / 임베딩 마이그레이션 반영)")
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph()

# ── 한눈에 보는 요약 ──────────────────────────────
H1("한 줄 요약")
P("매일 새벽 5시에 전 세계 AI 뉴스 사이트 22곳에서 기사를 긁어와, 같은 사건을 Gemini 임베딩으로 묶고, 공식 발표(Tier 0) → 중복 보도 많은 순으로 5개를 골라, 한국어·영어·일본어로 요약하고 음성까지 만들어 자동으로 웹사이트에 올리는 개인 뉴스 브리핑 서비스.")

doc.add_paragraph()
H1("내 컴퓨터에서도 실행됨")
BULLET("GitHub(웹사이트): 새벽 5시에 자동으로 빌드 → 웹사이트 갱신")
BULLET("내 컴퓨터: 아침 9시에 자동으로 창 띄우기 → Edge 앱 창으로 팝업")
BULLET("즉 새벽에 웹사이트가 먼저 업데이트되고, 내가 아침에 출근할 때 팝업이 뜨는 구조")

# ── 파일 역할 ─────────────────────────────
H1("파일은 딱 3개")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "파일"
hdr[1].text = "쉽게 말하면"
for name, desc in [
    ("signal_core.py", "진짜 엔진. 뉴스 긁기, 같은 사건 묶기, AI 필터, Gemini 요약, 음성 생성, HTML 만들기, 에러 처리까지 다 여기 들어있음"),
    ("ai_briefing.py", "내 PC에서 돌릴 때 쓰는 파일. 엔진을 돌린 뒤 Edge 창을 앱 모드로 팝업"),
    ("build.py", "GitHub에서 돌릴 때 쓰는 파일. 엔진을 돌린 뒤 결과를 public/ 폴더에 넣고 GitHub Pages에 배포. archive 폴더에 오늘 기사도 보존"),
]:
    row = tbl.add_row().cells
    row[0].text = name
    row[1].text = desc

doc.add_paragraph()

# ── 전체 흐름 ──────────────────────────────
H1("전체 흐름 — 새벽 5시에 일어나는 일 (순서대로)")

H2("1단계: 뉴스 긁기 (RSS 수집)")
P("뉴스 사이트 22곳에 접속해서 최근 24시간 동안 올라온 기사 제목과 요약문을 가져옴. 만약 24시간 안에 기사가 5개 미만이면 36시간으로 범위를 늘림. 죽은 피드는 실패로 기록해 상단에 경고 배너로 표시.")

P("수집하는 사이트 22곳:", bold=True)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "등급"
hdr[1].text = "매체"
hdr[2].text = "분류"
rows = [
    ("Tier 0", "OpenAI News", "공식 1차 소스 (ChatGPT/GPT)"),
    ("Tier 0", "Google Research", "공식 1차 소스 (응용/제품 연구)"),
    ("Tier 0", "Google DeepMind", "공식 1차 소스 (Gemma, AlphaFold)"),
    ("Tier 0", "Anthropic News", "공식 1차 소스 (Claude, 커뮤니티 피드)"),
    ("Tier 0", "NVIDIA Blog", "공식 1차 소스 (GPU/Physical AI)"),
    ("Tier 0", "Microsoft Research", "공식 1차 소스 (Copilot/Azure AI)"),
    ("Tier 1", "TechCrunch AI", "종합 테크 (속보)"),
    ("Tier 1", "The Verge AI", "종합 테크 (속보)"),
    ("Tier 1", "VentureBeat AI", "종합 테크 (엔터프라이즈)"),
    ("Tier 1", "The Decoder", "AI 전문 (심층 분석, 독일)"),
    ("Tier 1", "IEEE Spectrum AI", "기술 학회지 (IEEE 공식)"),
    ("Tier 1", "MIT Technology Review AI", "권위 저널리즘 (MIT)"),
    ("Tier 2", "MarkTechPost", "AI 전문 (인도, 논문/속보)"),
    ("Tier 2", "DailyAI", "AI 전문 (대량 속보)"),
    ("Tier 2", "Synced Review", "AI 전문 (중국 기반, 영어)"),
    ("Tier 2", "The Rundown AI", "AI 뉴스레터 (주간)"),
    ("Tier 2", "ZDNet AI", "실용 비교지"),
    ("Tier 2", "Simon Willison's Weblog", "AI 해설 개인 블로그"),
    ("Tier 2", "AI타임스 - AI기술", "한국어 AI (S1N24)"),
    ("Tier 2", "AI타임스 - AI산업", "한국어 AI (S1N3)"),
    ("Tier 2", "ITmedia AI+", "일본어 AI"),
    ("Tier 3", "Wired AI", "종합 테크 (문화/사회 앵글)"),
]
for a, b, c in rows:
    row = tbl.add_row().cells
    row[0].text = a
    row[1].text = b
    row[2].text = c

doc.add_paragraph()
P("등급이 뭘 뜻하냐면:", bold=True)
BULLET("Tier 0 — OpenAI/Google 같은 회사가 직접 발표하는 공식 블로그. 한 곳만 보도해도 최상위로 올라감")
BULLET("Tier 1~3 — 매체 신뢰도 순서. 중복 보도 수가 같을 때 순위 정하는 데 쓰임")
BULLET("Tier 3는 AI와 무관한 기사가 많이 섞여서 낮음")

# ── 매체 선정 이유 (22개 전부) ───────────────────
doc.add_page_break()
H1("매체 선정 이유 — 22개 전부 한 줄씩")
P("각 매체를 왜 넣었는지 구체적 근거. 원래 13개는 이 프로젝트 초기 사용자 본인 선정(일부 추정), 나머지는 임베딩 마이그레이션 후 추가.", bold=False)

H2("Tier 0 — 공식 1차 소스 (6개)")
BULLET("OpenAI News — ChatGPT/GPT 시리즈 제작사. AI 업계 단독 1위. 공식 발표가 업계에서 가장 파급력 큰 뉴스원")
BULLET("Google Research — Google의 응용/제품 연구 블로그. 교육 AI, Gemini 기반 제품, 사용자 경험 연구 등")
BULLET("Google DeepMind — Google Research와 별개 피드. Gemini/Gemma 본진 + AlphaFold/AlphaProteo 같은 과학 AI 발표처. 응용(Research) 대 프론티어(DeepMind) 역할 분담")
BULLET("Anthropic News — Claude 제작사. 공식 RSS 없어서 taobojlen/anthropic-rss-feed 커뮤니티 피드 사용 (매일 자동 갱신 확인). 발행 주 1~2회지만 나오면 무조건 중요")
BULLET("NVIDIA Blog — AI 하드웨어 사실상 독점. CUDA/GPU 로드맵, Physical AI/로보틱스 주도. AI 인프라 1차 소스")
BULLET("Microsoft Research — Copilot/Azure AI, Phi 모델 연구. OpenAI와 가장 깊은 전략 파트너라서 OpenAI 발표와 상호 보완")

H2("Tier 1 — 속보 + 권위 (6개)")
P("원래 있던 4개 (사용자 초기 선정):", bold=True)
BULLET("TechCrunch AI — 미국 테크 스타트업/M&A 속보 표준. 7일 교차율 24%로 다른 매체와 중복 높음 = 속보 빠름")
BULLET("The Verge AI — 소비자 관점 AI 제품/정책 뉴스. 7일 교차율 39%로 가장 높음 = 다른 매체들이 Verge를 받아씀")
BULLET("VentureBeat AI — 엔터프라이즈 AI + 투자/기업 분석에 강함. BtoB 앵글")
BULLET("The Decoder — 독일 기반 AI 전문지, 영어 발행. 속보보다 심층 분석 위주. 교차율 17%")
P("신규 추가 2개 (검증 완료):", bold=True)
BULLET("IEEE Spectrum AI — IEEE 공식 기술 학회지. 30 entries, 2026-04-13 최신. 학술적 권위 최상급 ('12 Graphs That Explain the State of AI in 2026' 같은 데이터 기반 심층 기사)")
BULLET("MIT Technology Review AI — MIT 기반 권위 저널리즘. 'Why opinion on AI is so divided' 같은 사회적·정책적 시각. 심층 기사 매일 발행")

H2("Tier 2 — AI 전문 매체 (9개)")
P("원래 있던 7개 (사용자 초기 선정):", bold=True)
BULLET("MarkTechPost — 인도 기반, AI 논문 해설 + 튜토리얼 + 속보. 매일 다수 기사로 물량 확보")
BULLET("DailyAI — AI 전문, 매일 발행. 대량 속보 커버리지")
BULLET("Synced Review — 중국 기반 AI 전문 매체 (영어 발행). 중국 AI 생태계 소식 강점")
BULLET("The Rundown AI — 주간 뉴스레터지만 7일 교차율 40%로 의외로 높음. 주간 큐레이션이라 중요 이슈가 모임")
BULLET("AI타임스 - AI기술 (S1N24) — 한국 최대 AI 전문지의 기술 섹션. 원래 allArticle 피드였는데 에너지/지역 기사가 섞여서 카테고리 분리해서 교체")
BULLET("AI타임스 - AI산업 (S1N3) — 같은 매체의 산업/기업 동향 섹션. AI기술과 상호 보완")
BULLET("ITmedia AI+ — 일본 최대 테크 매체의 AI 전용 섹션. 일본어 AI 뉴스의 표준 소스, 유일한 일본어 커버")
P("신규 추가 2개:", bold=True)
BULLET("ZDNet AI — 영미권 대형 테크 매체 AI 섹션. 'ChatGPT Plus vs Gemini Pro' 같은 실용적 비교 기사 강점. 소비자/개발자 관점")
BULLET("Simon Willison's Weblog — Django 공동 창시자, 현재 AI 해설로 가장 영향력 있는 독립 블로거 중 하나. 매일 업데이트, 다른 매체들도 자주 인용")

H2("Tier 3 — 종합 테크 (1개)")
BULLET("Wired AI — 문화/사회/윤리 관점 AI 기사 강점. AI 외 기사가 섞여서 낮은 Tier. 하지만 교차율 14%로 은근히 괜찮고 사회적 맥락 기사가 중요할 때 값어치 있음")

H2("제거한 2개 (7일 실제 교차율 데이터 기반)")
BULLET("Ars Technica — 7일 교차율 0% = 다른 매체와 중복 보도가 단 한 번도 없음. 종합 테크라 AI 외 기사도 많음 → 사실상 기여 제로")
BULLET("AI News (artificialintelligence-news.com) — 교차율 8% + 오늘 XML 파싱 에러 발생. 불안정 + 저성과")

H2("카테고리/언어 밸런스")
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Light Grid Accent 1"
hdr2 = tbl2.rows[0].cells
hdr2[0].text = "분류"
hdr2[1].text = "개수"
hdr2[2].text = "언어"
for a, b, c in [
    ("공식 1차 소스 (Tier 0)", "6", "EN"),
    ("속보 + 권위 (Tier 1)", "6", "EN"),
    ("AI 전문 매체 (Tier 2)", "9", "EN 6 / KR 2 / JP 1"),
    ("종합 테크 (Tier 3)", "1", "EN"),
    ("합계", "22", "EN 19 / KR 2 / JP 1"),
]:
    row = tbl2.add_row().cells
    row[0].text = a
    row[1].text = b
    row[2].text = c

doc.add_paragraph()
P("언어 밸런스는 영어 중심. 한국/일본 매체 확장하려면 신규 RSS 소스 검증 필요.")
P("주제 밸런스: 속보(Tier 1 종합 테크) + 심층(IEEE Spectrum, MIT Tech Review) + 해설(Simon Willison) + 공식 1차(Tier 0) + 실용 비교(ZDNet) → 대부분 각도 커버됨.")

H2("2단계: 같은 사건 묶기 (Gemini 임베딩)")
P("핵심 변경점. 예전엔 TF-IDF(단어 겹침)로 묶었는데, 같은 사건인데 매체마다 표현이 달라서 못 묶는 경우가 많았음. 특히 한국어·일본어는 영어와 단어가 아예 안 겹쳐서 무용지물.")

P("지금은 Gemini 임베딩 모델(gemini-embedding-001)을 써서 '의미 기반'으로 묶음:", bold=True)
BULLET("각 기사의 제목+요약을 Gemini가 숫자 벡터(3072차원)로 변환")
BULLET("두 벡터가 방향이 비슷하면 같은 의미라고 판단 (코사인 유사도)")
BULLET("기준점: 0.85 이상이면 같은 사건으로 묶음")
BULLET("같은 매체끼리는 묶지 않음 (중복 보도 숫자가 의미 없어져서)")

P("쉽게 비유하면:", bold=True)
P("\"GPT-5 발표\"와 \"OpenAI, 차세대 모델 공개\"가 단어는 다르지만 의미가 거의 같다 → 벡터가 비슷한 방향을 가리킴 → 0.85 넘음 → 같은 사건으로 묶임. 영어·한국어·일본어를 넘나들어도 의미가 같으면 묶이는 게 핵심.")

P("실제 테스트 결과:", bold=True)
BULLET("Sam Altman 공격 사건이 The Verge(영어) + The Decoder(영어) + ITmedia(일본어) + AI타임스(한국어) 4곳에서 보도 → 전부 하나로 묶임")
BULLET("Apple 스마트 안경, OpenAI Spud 메모, 유니트리 로봇 기사도 다국어 매칭 성공")

P("기술적 설정:", bold=True)
BULLET("배치 처리: 한 번에 100개씩 보냄 (Gemini API 요청당 20,000 토큰 제한 때문)")
BULLET("배치 사이 20초 대기 (분당 요청 수 제한에 걸리지 않게)")
BULLET("660개 기사 처리 시 대략 2~3분 소요")
BULLET("무료 티어 한도(RPM 100, TPM 1000만)의 5% 미만 사용")

H2("2.5단계: AI 아닌 뉴스 거르기")
P("AI타임스 같은 한국어 매체는 에너지·지역 뉴스까지 섞여 들어오기 쉬움. 그래서 단순한 키워드 검사로 한 번 더 걸러냄.")

P("화이트리스트(= 이 단어가 있으면 AI 뉴스로 통과):", bold=True)
BULLET("영어: LLM, GPT, transformer, RAG, agent, diffusion, robot, GPU, prompt, chatbot 등")
BULLET("한국어: 인공지능, 머신러닝, 딥러닝, 생성형, 언어모델, 클로드, 제미나이, 라마, 챗봇, 로봇, 자율주행, 에이전트, 프롬프트, 파인튜닝 등")
BULLET("일본어: 人工知能, 機械学習, 深層学習, 生成AI, 大規模言語モデル, クロード, ジェミニ, ロボット 등")

P("블랙리스트(= 이 단어가 있는데 화이트리스트엔 안 걸리면 제외):", bold=True)
BULLET("영어: war, military, weapon, election, tariff, crime, kill 등")
BULLET("한국어: 재생에너지, 태양광, 풍력, 가정용 배터리, 지역축제, 추경 예산, 선거, 전쟁 등")
BULLET("일본어: エネルギー転換, 再生可能エネルギー, 戦争, 選挙 등")

P("판정 규칙:", bold=True)
BULLET("블랙만 걸리고 화이트는 안 걸림 → 확실히 AI 무관 → 제외")
BULLET("블랙도 있고 화이트도 있음 → 애매함 → 통과시키고 Gemini에게 맡김")
BULLET("블랙 없음 → 통과")

H2("3단계: Gemini로 한/영/일 요약 만들기")
P("여기는 매일 API를 딱 한 번만 호출. 상위 8개 토픽을 통째로 Gemini에 넘겨서, 그중 5개를 골라 3개국어 요약을 한 번에 받음.")

P("쓰는 모델:", bold=True)
BULLET("gemini-2.5-flash (무료 티어)")
BULLET("요청당 재시도 최대 3회")
BULLET("재시도 간격 10 → 20 → 40초 (서버 과부하 503 대응)")

doc.add_page_break()
H2("Gemini에 실제로 보내는 프롬프트 원문")
P("코드에서 f-string으로 조립돼서 이 문장이 Gemini에 전달됨. {today}는 오늘 날짜, {articles_text}는 상위 8개 토픽 정보, {NEWS_COUNT}는 5가 들어감.")

prompt_text = '''오늘은 {today}입니다.
아래는 AI 뉴스를 중복 보도 수 기준으로 정렬한 토픽 목록입니다.
중복 보도 수는 이미 코드에서 계산된 정확한 값입니다.

{articles_text}

## 규칙 (엄격히 준수)
1. AI 기술/제품/연구와 직접 관련 없는 토픽은 무조건 제외하세요. AI 회사가 언급되더라도 내용이 정치, 군사, 외교, 부동산, 데이터센터 물리적 위협 등이면 제외합니다.
2. 남은 토픽 중 상위 {NEWS_COUNT}개를 선택하세요 (중복 보도 수가 많은 순).
3. 각 토픽의 "원문 요약"에 있는 정보만 사용하세요. 원문에 없는 내용을 추가하거나 추측하지 마세요.
4. 한국어로 번역/요약하되, 사실관계를 변경하지 마세요.
5. "topic_index"는 위 목록의 [토픽 N]에서 N을 그대로 사용하세요.

다른 텍스트 없이 JSON 배열만 출력하세요.

[
  {
    "topic_index": 토픽 번호 (정수),
    "tag_kr": "카테고리 한국어 (LLM, 이미지AI, 로보틱스, 규제, 기업동향, 연구 중 하나)",
    "tag_en": "카테고리 영어 (LLM, Image AI, Robotics, Regulation, Industry, Research 중 하나)",
    "tag_jp": "카테고리 일본어 (LLM, 画像AI, ロボティクス, 規制, 企業動向, 研究 중 하나)",
    "title_kr": "한국어 제목 (20자 이내)",
    "title_en": "영어 제목 (20자 이내)",
    "title_jp": "일본어 제목 (20자 이내)",
    "summary_kr": "한국어 핵심 요약 (2~3문장)",
    "summary_en": "영어 핵심 요약 (2~3문장)",
    "summary_jp": "일본어 핵심 요약 (2~3문장)",
    "why_kr": "왜 중요한가 (한국어 1문장)",
    "why_en": "왜 중요한가 (영어 1문장)",
    "why_jp": "왜 중요한가 (일본어 1문장)"
  }
]'''
for line in prompt_text.split("\n"):
    CODE(line if line else " ")

P("articles_text 안은 이렇게 조립됨 (상위 8개 토픽 반복):", bold=True)
for line in [
    "[토픽 1] (중복 보도: 4개 매체)",
    "대표 제목: (매체 신뢰도 최상위 기사의 원문 제목)",
    "원문 요약: (RSS 요약 300자 이내)",
    "보도 매체: TechCrunch, The Verge, VentureBeat, The Decoder",
]:
    CODE(line)

H3("이 프롬프트 설계 의도 (왜 이렇게 썼는가)")
BULLET("중복 보도 수를 코드에서 먼저 계산하고 프롬프트에 고정값으로 박음 → Gemini가 숫자를 바꾸거나 추측 못 하게 막음")
BULLET("규칙 1: 'AI 회사가 언급되더라도 정치·군사면 제외' → 코드 블랙리스트가 놓친 애매한 케이스를 잡는 최후 방어선")
BULLET("규칙 3: '원문에 있는 정보만 사용' → Gemini가 없는 사실을 만들어내는 환각을 줄임")
BULLET("규칙 4: 한국어로 번역하되 사실관계 그대로 → 의역 금지")
BULLET("topic_index를 1번부터 받고 코드에서 0번부터로 변환 → Gemini가 쉬운 번호 체계 사용")
BULLET("JSON 배열만 출력하라고 명령 → 파싱 안정성")
BULLET("링크, 매체 수, 파비콘 아이콘은 프롬프트에 안 포함 → 어차피 코드에서 원본 값으로 덮어씀 → Gemini가 환각 만들 여지 차단")

H3("Gemini 응답을 받으면 코드가 하는 일")
BULLET("응답 텍스트에서 ```json 같은 감싸는 기호 제거")
BULLET("JSON 파싱")
BULLET("topic_index로 원본 토픽 다시 매칭")
BULLET("링크, 매체 수 텍스트(한/영/일), 원본 제목, 발행일 → 전부 코드에서 덮어씀 (Gemini 말 안 믿음)")
BULLET("파비콘 도메인: 제목과 본문에서 회사명 찾아서 해당 회사 도메인 매핑 (OpenAI→openai.com, Meta→meta.com 등)")
BULLET("검증 통과한 뉴스가 5개 미만이면 재시도")
BULLET("⭐ 최종 정렬 복원: Gemini가 프롬프트의 '중복 보도 수 많은 순' 규칙 때문에 Tier 0을 뒤로 밀어내는 경우가 있어, 받은 결과를 topic_index 순으로 다시 정렬해서 cluster_articles의 원래 정렬(Tier 0 우선)을 복원함")

H2("4단계: 음성 만들기 (Edge TTS)")
P("각 뉴스 카드에 재생 버튼이 있어서 클릭하면 한/영/일 세 가지 음성으로 뉴스를 읽어줌. Microsoft Edge의 TTS(무료)를 씀.")

BULLET("한국어: ko-KR-SunHiNeural")
BULLET("영어: en-US-AriaNeural")
BULLET("일본어: ja-JP-NanamiNeural")
BULLET("속도는 기본보다 5% 빠르게 (+5%)")
BULLET("각 뉴스마다 '1번째 뉴스입니다', 'News number 1', '1番目のニュースです' 식 인트로 자동 추가")
BULLET("개별 뉴스 음성 + 전체 이어 재생 음성 둘 다 만듦")
BULLET("로컬 실행: base64 data URI로 HTML에 통째로 박아서 파일 하나로 완결")
BULLET("배포 실행: public/audio/ 폴더에 mp3 파일로 저장")

H2("5단계: HTML 웹페이지 만들기")
BULLET("다크 테마(검정 배경에 보라/인디고 포인트)")
BULLET("첫 카드는 'featured'로 강조 (크고 눈에 띄게)")
BULLET("Spotify 스타일 하단 플레이어 바 (재생/일시정지/이전/다음)")
BULLET("오로라 Canvas 배경 애니메이션")
BULLET("KR/EN/JP 언어 전환 버튼 — JavaScript setLang() 함수로 바로 바뀜")
BULLET("각 카드: 번호, 카테고리 태그, 매체 수, 회사 파비콘, 제목, 요약, 왜 중요한가(WHY IT MATTERS), 재생 버튼, 원문 링크(3개국어 번역됨), 시각")
BULLET("피드 실패 시 상단에 주황/빨강 경고 배너")

H2("6단계: 배포 (build.py 경로만 해당)")
BULLET("public/index.html — 오늘 브리핑 메인 페이지")
BULLET("public/archive/YYYY-MM-DD.html — 오늘 분을 아카이브에 복사")
BULLET("archive/ 폴더도 git에 커밋 → 영구 보존")
BULLET("metrics/YYYY-MM-DD.json — 매체별 교차 보도율 통계")
BULLET("GitHub Pages로 자동 배포")
BULLET("Discord 웹훅이 설정돼 있으면 헤드라인 알림 전송 (현재 미설정)")

# ── 자동 갱신 & 자동 학습 ─────────────────────────
doc.add_page_break()
H1("매체 신뢰도 자동 학습")
P("매일 각 매체가 얼마나 다른 매체와 같은 사건을 보도했는지(교차 보도율) 기록함. 7일치가 쌓이면 자동으로 등급을 재조정:")

BULLET("교차율 25% 이상 → Tier 1로 승급 (가중치 3)")
BULLET("10~25% → Tier 2 (가중치 2)")
BULLET("10% 미만 → Tier 3 (가중치 1)")
BULLET("Tier 0(OpenAI/Google 등 공식 소스)는 자동 학습에서 제외 → 수동 유지")

P("이 등급은 나중에 같은 매체 수로 묶인 토픽 중 어느 걸 위에 올릴지, 그리고 대표 기사를 누구 걸 쓸지 결정할 때 사용됨.")

# ── 정렬 & Tier 0 특별 대우 ───────────────────────
H1("뉴스 순서 정하는 법")
P("클러스터링 후 상위 토픽을 뽑을 때 다음 순서로 정렬:")

BULLET("1순위: Tier 0 매체가 보도했는가? (있으면 무조건 위로)")
BULLET("2순위: 매체 수가 몇 개인가? (많을수록 위)")
BULLET("3순위: 보도한 매체 중 가장 신뢰도 높은 등급은? (높을수록 위)")

P("즉 OpenAI가 자기 블로그에 발표한 뉴스는 다른 매체가 보도 안 해도 1위. 그다음부터는 중복 보도 수가 많은 순.")

# ── 에러 처리 ──────────────────────────────
H1("에러 나면 어떻게 되는가")
P("빌드가 실패하면 빨간 '빌드 실패' 페이지를 대신 올림. 에러 메시지가 상세하게 표시되고, 아카이브 링크는 살아있어서 어제 분은 계속 볼 수 있음.")

P("실패로 이어지는 주요 원인:", bold=True)
BULLET("GEMINI_API_KEY 미설정")
BULLET("RSS 피드에서 24~36시간 안에 단 한 건도 못 가져온 경우 (전체 다운)")
BULLET("Gemini 큐레이션 3회 전부 실패 (서버 과부하 503 등)")
BULLET("Gemini가 검증 통과 뉴스 5개를 못 채움 (데이터 부족)")

P("실패 안 하고 그냥 경고만 뜨는 경우:", bold=True)
BULLET("일부 피드 실패 → 상단 배너(주황/빨강)로 표시, 빌드는 계속")
BULLET("TTS 일부 실패 → 해당 음성만 빈 상태, 나머지 정상")
BULLET("Discord 알림 실패 → 조용히 스킵")

# ── 자동 실행 ──────────────────────────────
H1("자동 실행 스케줄")
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "실행 주체"
hdr[1].text = "시각 (KST)"
hdr[2].text = "하는 일"
for a, b, c in [
    ("GitHub Actions", "매일 05:00", "웹사이트(GitHub Pages) 자동 갱신"),
    ("Windows 작업 스케줄러", "매일 09:00", "내 PC에 Edge 앱 창으로 팝업"),
]:
    row = tbl.add_row().cells
    row[0].text = a
    row[1].text = b
    row[2].text = c

doc.add_paragraph()
P("즉 새벽 5시에 웹사이트가 먼저 갱신되고, 4시간 뒤 아침 9시에 내 PC가 팝업으로 보여주는 구조.")

# ── 환경 변수 ──────────────────────────────
H1("환경 변수")
BULLET("GEMINI_API_KEY — 필수. 없으면 에러 페이지 생성하고 종료")
BULLET("DISCORD_WEBHOOK — 선택. 있으면 빌드 결과 알림 전송")
BULLET("PAGES_URL — 기본값 https://davidking981218-rgb.github.io/the-signal/")

# ── 이번 개정에서 바뀐 것 ────────────────────
doc.add_page_break()
H1("이번 개정에서 바뀐 것 (마이그레이션 내역)")

P("1. SDK 전면 교체 — google.generativeai는 Google이 공식 deprecate했음. 새 SDK google.genai로 전면 이전.", bold=True)
BULLET("기존: genai.configure(api_key=...), genai.GenerativeModel(...).generate_content(prompt)")
BULLET("신규: client = genai.Client(api_key=...); client.models.generate_content(model=..., contents=...)")

P("2. 클러스터링: TF-IDF → Gemini 임베딩", bold=True)
BULLET("기존 문제: 단어 겹침 기반이라 다국어(한/일) 매칭 불가, 영어도 표현 다르면 못 묶음")
BULLET("신규: gemini-embedding-001로 의미 벡터 생성, 코사인 유사도 0.85 이상이면 묶음")
BULLET("배치 100개씩, 배치 사이 20초 대기 (rate limit 여유 확보)")
BULLET("sklearn 의존성 완전 제거")

P("3. 피드 확장: 13개 → 22개", bold=True)
BULLET("Tier 0 공식 소스 6개 추가: OpenAI News, Google Research, Google DeepMind, Anthropic News, NVIDIA Blog, Microsoft Research")
BULLET("(Hugging Face는 일시 추가했다가 '커뮤니티 튜토리얼 중심이라 뉴스성 부족'으로 제거)")
BULLET("Tier 1 확장: IEEE Spectrum AI, MIT Technology Review AI 추가 (권위 기술 저널)")
BULLET("Tier 2 확장: ZDNet AI, Simon Willison's Weblog 추가 (실용지 + AI 해설 블로그)")
BULLET("AI타임스 allArticle → AI기술(S1N24) + AI산업(S1N3) 카테고리 분리 (에너지·지역 기사 오염 해결)")
BULLET("Anthropic은 공식 RSS 없어서 커뮤니티 피드(taobojlen/anthropic-rss-feed) 사용 — 매일 자동 갱신 확인됨")
BULLET("제거: Ars Technica (7일 교차율 0%), AI News (교차율 8% + XML 파싱 에러)")

P("4. 정렬 로직 변경", bold=True)
BULLET("기존: 매체 수 → 매체 신뢰도")
BULLET("신규: Tier 0 여부 → 매체 수 → 매체 신뢰도 (공식 발표는 단독이라도 최상위)")
BULLET("⭐ curate_with_gemini 후처리에서도 topic_index 순으로 재정렬 → Gemini가 순서를 바꿔도 Tier 0 우선이 최종 출력에 보존됨")

P("5. AI 관련성 필터에 한국어/일본어 확장", bold=True)
BULLET("기존: 영어 정규식만 있어서 한/일 기사는 판별 불가능 → 그냥 통과")
BULLET("신규: 한/일 substring 리스트 추가 (인공지능, 人工知能, 생성AI, LLM 등 화이트 + 재생에너지, 선거 등 블랙)")

P("6. 매체 신뢰도 자동 학습에서 Tier 0 보호", bold=True)
BULLET("7일치 통계 쌓여 자동 갱신될 때 Tier 0 매체는 건너뜀")

P("7. Gemini 재시도 간격 강화", bold=True)
BULLET("기존: 1 → 2 → 4초 (합 7초, 503 과부하 풀리기 전에 소진)")
BULLET("신규: 10 → 20 → 40초 (합 70초, 일시 과부하 회복 시간 확보)")

P("8. 원문 보기 버튼 3개국어 번역", bold=True)
BULLET("KR: 원문 보기 ↗")
BULLET("EN: Read original ↗")
BULLET("JP: 原文を見る ↗")

P("9. GitHub Actions 의존성 업데이트", bold=True)
BULLET("기존: feedparser + google-generativeai + scikit-learn + edge-tts")
BULLET("신규: feedparser + google-genai + numpy + edge-tts")

# ── 아직 남은 이슈 ─────────────────────────
H1("아직 남은 한계")
BULLET("Gemini 서버 과부하(503) — 인프라 문제라 코드로 해결 불가, 재시도 간격 늘리는 것이 한계")
BULLET("모바일 대응 미검증")
BULLET("언어 설정이 KR/EN/JP 3개로 여전히 여러 곳에 하드코딩되어 있음 (새 언어 추가 시 최소 6곳 수정)")
BULLET("트위터, Reddit, Hacker News, arXiv 같은 비-RSS 소스는 수집하지 않음")
BULLET("개인화 필터 없음 (모두 같은 상위 5개를 봄)")
BULLET("과거 브리핑 키워드 검색 기능 없음")
BULLET("아카이브가 git 커밋 방식이라 장기적으로 레포가 커짐 (SQLite 전환 고려 대상)")

doc.save(r"C:\Users\user\ai_briefing\THE_SIGNAL_설명서.docx")
print("OK")
