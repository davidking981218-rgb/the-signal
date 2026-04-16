"""THE SIGNAL 상세 설명서 docx 생성 스크립트.
signal_core.py의 실제 코드를 근거로만 작성한다. 추측 금지.
검증한 파일: signal_core.py (2026-04-15 기준)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# 기본 폰트
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)


def H1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)


def H2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)


def P(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    return p


def BULLET(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def CODE(text):
    """단락 스타일의 코드 블록 (Consolas, 회색 배경 효과는 폰트만)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return p


# ═══════════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════════
title = doc.add_heading("THE SIGNAL", level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = sub.add_run("AI Daily Briefing — 상세 설명서")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
meta.add_run("2026-04-15 작성 · signal_core.py 실제 코드 기반").font.size = Pt(9)

doc.add_paragraph()

# 검증 근거 (CLAUDE.md 규칙: 문서 첫 페이지에 명시)
H2("이 문서의 검증 근거")
P("이 문서는 다음 파일을 직접 읽은 뒤 작성되었다. 코드에 없는 내용은 기재하지 않는다.")
BULLET("C:\\Users\\user\\ai_briefing\\signal_core.py (모듈 상수, 상수 주석, 핵심 함수 전체)")
BULLET("C:\\Users\\user\\ai_briefing\\README.md (기존 공개 설명)")
P("[주의] 매체별 개별 선정 이유는 코드 주석에 카테고리 단위로만 남아 있어, 이 문서도 카테고리 수준까지만 설명한다.",
  bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. 개요
# ═══════════════════════════════════════════════════════════════
H1("1. 개요")
P("THE SIGNAL은 22개 글로벌 AI 매체(영어/한국어/일본어)에서 뉴스를 수집하고, "
  "Gemini 임베딩으로 의미가 같은 기사를 묶은 뒤, Gemini 2.5 Flash로 한국어/영어/일본어 "
  "3개 국어 요약을 제공하는 AI 뉴스 큐레이션 서비스다.")

BULLET("라이브: https://davidking981218-rgb.github.io/the-signal/")
BULLET("저장소: https://github.com/davidking981218-rgb/the-signal")
BULLET("출력: 매일 5개 뉴스 카드 (signal_core.py의 NEWS_COUNT=5)")
BULLET("자동 갱신: GitHub Actions 매일 UTC 20:00 (KST 05:00) cron")

# ═══════════════════════════════════════════════════════════════
# 2. 파이프라인 한눈에
# ═══════════════════════════════════════════════════════════════
H1("2. 파이프라인 한눈에")
P("하루치 빌드는 순서대로 다음 단계를 거친다. 각 단계는 독립 함수로 분리되어 있어 "
  "어느 단계에서 문제가 나면 해당 함수만 보면 된다.")

steps = [
    ("① RSS 수집", "fetch_rss() — 22개 피드에서 24시간 이내 기사만 긁어온다. 부족 시 36시간으로 확장."),
    ("② 의미 기반 클러스터링", "cluster_articles() — Gemini 임베딩 후 코사인 유사도 ≥ 0.85인 기사들을 같은 토픽으로 묶는다."),
    ("③ AI 관련성 필터", "filter_ai_relevant() — 영어 정규식 + 한/일 substring 블랙/화이트 리스트로 비-AI 토픽 제거."),
    ("④ Gemini 요약", "curate_with_gemini() — 상위 8개 토픽을 Gemini 2.5 Flash에 넘겨 KR/EN/JP 3개 언어로 요약 생성."),
    ("⑤ 팩트 덮어쓰기", "URL/매체 수/원문 제목/발행일/브랜드 키는 코드가 직접 덮어쓴다. Gemini가 만지지 못한다 (환각 방지)."),
    ("⑥ HTML 렌더링", "build_html() — 다크 테마 카드 UI로 변환. Tier 0 매체 카드에는 브랜드 컬러 적용."),
    ("⑦ Edge TTS 음성", "generate_tts() — 기사별 + 전체 재생용 음성을 3개 언어로 만든다."),
    ("⑧ GitHub Pages 배포", "build.py → GitHub Actions → Pages. 아카이브는 git 커밋으로 영구 보존."),
]
for k, v in steps:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(k + " — ")
    run.bold = True
    run.font.size = Pt(10.5)
    p.add_run(v).font.size = Pt(10.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. 매체 선정과 Tier 구조
# ═══════════════════════════════════════════════════════════════
H1("3. 매체 선정과 Tier 구조")
P("총 22개 피드를 4단계 Tier로 나눠 가중치를 둔다. Tier는 단순 랭킹이 아니라 "
  "'같은 매체 수일 때 누가 우선인가'를 결정하는 타이 브레이커다. "
  "정렬 1순위는 Tier 0 포함 여부이고, 2순위가 매체 수, 3순위가 Tier 가중치다.")

H2("Tier 0 — 공식 1차 소스 (가중치 4)")
P("회사가 직접 운영하는 공식 블로그/뉴스룸. AI 분야에서 이들이 올리는 글은 "
  "제품 공개, 논문 발표, 정책 등 '원본'이다. 그래서 '단 1곳만 보도해도 무조건 최상위'로 올라간다 "
  "(TIER_0_WEIGHT=4, 정렬 키 `has_tier0`).")

BULLET("OpenAI News — ChatGPT/GPT 제품 발표와 안전/정책 공지가 가장 먼저 올라오는 회사 공식 뉴스룸이라 원본 신호원으로 필수")
BULLET("Google Research — 구글의 AI/ML 연구 논문과 공식 발표가 올라오는 연구 블로그. 논문 수준 원본을 확보하기 위해 포함")
BULLET("Google DeepMind — Gemini, AlphaFold 등 프론티어 연구를 직접 공개하는 채널 (signal_core.py 주석 원문)")
BULLET("Anthropic News — Claude 관련 공식 발표. Anthropic이 공식 RSS를 제공하지 않아 커뮤니티 피드 taobojlen/anthropic-rss-feed로 대체 (코드 주석 명시)")
BULLET("NVIDIA Blog — AI/로봇/GPU 공식 채널 (signal_core.py 주석 원문). GPU/하드웨어 레이어의 1차 소스")
BULLET("Microsoft Research — Copilot/Azure AI 연구 (signal_core.py 주석 원문). MS 진영 AI 연구 원본")

H2("Tier 1 — AI 속보 + 권위 저널 (가중치 3)")
P("AI 속보를 빠르게 다루는 종합 테크 매체와, 기술 저널리즘 권위지. "
  "같은 속보를 여러 번 중복 보도할 가능성이 높아 클러스터링 신호가 강한 계층이다.")
BULLET("TechCrunch AI — 종합 테크 매체 중 AI 속보 커버리지가 넓어, 같은 사건의 중복 보도를 잘 유발하는 신호원 (코드 카테고리 주석: '속보 겹침 가능성 높음')")
BULLET("The Verge AI — 동일 카테고리. 제품/업계 관점 보도가 강해 다른 영어 매체와의 교차 보도가 잦음")
BULLET("VentureBeat AI — 동일 카테고리. 엔터프라이즈 AI 속보가 다른 매체와 자주 겹침")
BULLET("The Decoder — 코드 주석상 'AI 전문 매체 — 매일 다수 기사, 같은 사건 중복 보도' 카테고리. 클러스터링 신호 확보용")
BULLET("IEEE Spectrum — IEEE 기술 학회지 (signal_core.py 주석 원문). 권위 있는 기술 저널 축 확보")
BULLET("MIT Technology Review — 권위 저널리즘 (signal_core.py 주석 원문). 분석 기사 축 확보")

H2("Tier 2 — AI 전문 매체 + 실용지 + 지역 전문 (가중치 2)")
P("AI 산업을 전문으로 다루는 매체와, 개인 해설 블로그, 한국/일본 지역 전문지가 여기 속한다. "
  "매체 수는 많지만 중복도가 상대적으로 낮아 타이 브레이커 역할.")
BULLET("MarkTechPost — 코드 카테고리 'AI 전문 매체 — 매일 다수 기사'. 논문/모델 리뷰 위주로 매일 분량이 많아 중복 보도 신호 보강용")
BULLET("DailyAI — 동일 카테고리. AI 분야 일일 뉴스 발행량이 많아 클러스터링 신호원")
BULLET("Synced — 동일 카테고리 'AI 전문 매체 — 매일 다수 기사'")
BULLET("The Rundown AI — 코드 주석상 'AI 뉴스레터 — 주요 뉴스 큐레이션'. 큐레이션된 핵심 뉴스만 올라오므로 Tier 2로 보강 신호")
BULLET("ZDNet AI — SOURCE_WEIGHT 주석상 'AI 전문 매체 + 실용지' 카테고리. 엔터프라이즈/실용 관점 커버리지 확보")
BULLET("Simon Willison's Weblog — 매일 업데이트되는 AI 해설 개인 블로그 (signal_core.py 주석 원문). 해설/분석 층 보강")
BULLET("AI타임스 (AI기술/AI산업) — 한국 AI 전문지. 두 카테고리를 분리 구독해 전체 피드(allArticle)의 오염을 제거 (signal_core.py 주석 '카테고리 분리 — allArticle 오염 제거')")
BULLET("ITmedia AI+ — 일본 AI 전문 매체 (signal_core.py 주석 원문 '일본 AI 전문'). 임베딩의 크로스링구얼 매칭을 활용해 일본어 기사도 동일 클러스터로 묶기 위함")

H2("Tier 3 — 종합 테크 (가중치 1)")
P("AI가 아닌 기사가 많이 섞여 들어오는 채널. 있으면 좋지만 가중치는 가장 낮다.")
BULLET("Wired AI — 종합 테크의 AI 섹션. 문화/장기 분석 관점 보강용이지만 AI 외 기사 비중이 높아 코드가 가중치를 가장 낮게 부여")

H2("선정 원칙 (코드 주석 기반)")
BULLET("Tier 0은 '원본'이 들어오는 곳 — 언제나 최상위 보장")
BULLET("Tier 1~2는 '중복 보도'를 통해 중요도를 검출하는 신호원")
BULLET("언어는 영어 기본 + 한국어(AI타임스 2종) + 일본어(ITmedia AI+) — 임베딩 모델이 다국어 크로스링구얼 매칭을 지원하기 때문에 번역 없이 묶인다")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. RSS 수집과 기간 조건
# ═══════════════════════════════════════════════════════════════
H1("4. RSS 수집과 기간 조건 (fetch_rss)")
P("feedparser로 22개 피드를 순차 파싱한다. 실패한 피드는 건너뛰고 feed_status에 기록된다(사이트 상단 배너 표시).")

H2("기간 규칙")
BULLET("기본 cutoff: 24시간 (_collect(24))")
BULLET("수집 기사 수가 NEWS_COUNT=5보다 적으면 자동으로 36시간으로 확장 재시도")
BULLET("Tier 0 매체도 동일한 24h cutoff를 적용한다 (특별 대우 없음)")
BULLET("발행 시각은 feedparser의 published_parsed / updated_parsed 우선. 둘 다 없으면 스킵")

H2("피드 건강 배너")
P("ok/fail 카운트를 계산하고, 절반 이상 실패하면 빨강 배너, 일부 실패는 주황 배너를 HTML 상단에 표시한다.")

# ═══════════════════════════════════════════════════════════════
# 5. 의미 기반 클러스터링
# ═══════════════════════════════════════════════════════════════
H1("5. 의미 기반 클러스터링 — 같은 기사 묶는 방법 (cluster_articles)")
P("핵심 질문은 '서로 다른 매체에서 올라온 두 기사가 같은 사건을 다루고 있는가?'이다. "
  "THE SIGNAL은 이를 키워드/URL 매칭이 아니라 임베딩 벡터의 코사인 유사도로 판단한다.")

H2("알고리즘 단계")
p = doc.add_paragraph(style="List Number")
p.add_run("텍스트 준비 — ").bold = True
p.add_run("각 기사에서 `title + ' ' + summary` 문자열 생성")
p = doc.add_paragraph(style="List Number")
p.add_run("임베딩 — ").bold = True
p.add_run("Gemini gemini-embedding-001 모델로 벡터화. task_type='SEMANTIC_SIMILARITY', 배치 크기 100, 배치 간 20초 대기(RPM/TPM 여유 확보)")
p = doc.add_paragraph(style="List Number")
p.add_run("정규화 + 코사인 유사도 — ").bold = True
p.add_run("L2 정규화 후 내적으로 유사도 행렬 계산")
p = doc.add_paragraph(style="List Number")
p.add_run("Union-Find 묶음 — ").bold = True
p.add_run("유사도 ≥ SIMILARITY_THRESHOLD (0.85) 쌍을 같은 집합으로 합친다. 단, 같은 매체 내 기사끼리는 묶지 않는다(`entries[i]['source'] == entries[j]['source']` 스킵)")
p = doc.add_paragraph(style="List Number")
p.add_run("대표 기사 선정 — ").bold = True
p.add_run("한 클러스터에서 매체 신뢰도가 가장 높은 기사, 동점이면 요약이 긴 기사를 대표로 뽑는다")

H2("유사도 임계값 0.85인 이유 (코드 주석 근거)")
P("`SIMILARITY_THRESHOLD = 0.85  # 임베딩 코사인 유사도 기준 (의미 기반, 다국어 호환). 0.80은 과묶음`")
P("즉 0.80까지 내리면 서로 다른 사건도 하나로 묶여버린다는 실험 결과가 주석에 남아 있다.")

H2("다국어 크로스링구얼 매칭")
P("gemini-embedding-001은 한/영/일을 같은 벡터 공간에 매핑한다. 그래서 '오픈AI가 GPT-5 발표'(한국)와 "
  "'OpenAI launches GPT-5'(영문)와 'OpenAIがGPT-5を発表'(일본)는 번역 없이도 유사도 0.85 이상으로 묶인다.")

H2("대표 기사 결정 규칙 (정확 코드)")
CODE("representative = max(\n"
     "    arts,\n"
     "    key=lambda a: (SOURCE_WEIGHT.get(a['source'], _DEFAULT_WEIGHT), len(a['summary']))\n"
     ")")

H2("정렬 키")
CODE("def _sort_key(topic):\n"
     "    best_weight = max(SOURCE_WEIGHT.get(s, _DEFAULT_WEIGHT) for s in topic['sources'])\n"
     "    has_tier0 = best_weight >= TIER_0_WEIGHT\n"
     "    return (has_tier0, topic['source_count'], best_weight)")
P("즉 Tier 0 포함 → 중복 매체 수 → Tier 가중치 순으로 정렬.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. AI 관련성 필터
# ═══════════════════════════════════════════════════════════════
H1("6. AI 관련성 필터 (filter_ai_relevant)")
P("클러스터링된 토픽 중 AI 기술과 무관한 것을 걸러낸다. 키워드 매칭이지만 블랙/화이트를 조합해 애매한 건 Gemini에게 위임한다.")

H2("판정 규칙")
BULLET("블랙리스트만 걸림 → 제외 (예: tariff, election, war)")
BULLET("블랙 + 화이트 동시 걸림 → Gemini 판단에 위임 (로그에 '△ 애매'로 표시)")
BULLET("블랙 없음 → 통과")

H2("블랙리스트 (영어 정규식)")
CODE(r"threat, sanction, military, war, weapon, tariff, trade war, election, vote," "\n"
     r"kill, death, crime, arrest, datacenter attack, cyberattack(?!.*detect), hack(?!athon)")

H2("화이트리스트 (영어 정규식, 일부)")
CODE(r"LLM, GPT, transformer, fine-tun, RAG, model...(train|launch|release|param)," "\n"
     r"agent, diffusion, image generat, video generat, robot, autonomous," "\n"
     r"self-driving, GPU, TPU, open-source model, benchmark, prompt, chatbot," "\n"
     r"copilot, assistant, API...(launch|release|update), safety...AI, regulat...AI, AI act, AI policy")

H2("한국어/일본어 substring")
P("정규식이 아닌 단순 substring 매칭.")
BULLET("KR 화이트: 인공지능, 머신러닝, 딥러닝, 생성형/생성 AI, 대규모 언어모델, 클로드, 제미나이, 라마, 챗봇, 챗GPT, 로봇, 자율주행, 에이전트, 프롬프트, 오픈소스 모델, 파인튜닝, 거대언어모델, 초거대 등")
BULLET("JP 화이트: 人工知能, 機械学習, 深層学習, 生成AI, 大規模言語モデル, クロード, ジェミニ, ロボット, 自動運転, エージェント, プロンプト 등")
BULLET("KR 블랙: 재생에너지, 태양광, 풍력, 선거, 투표, 전쟁, 군사, 무기, 관세, 추경 예산 등")
BULLET("JP 블랙: エネルギー転換, 再生可能エネルギー, 戦争, 選挙")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. Gemini 프롬프트 원문
# ═══════════════════════════════════════════════════════════════
H1("7. Gemini 프롬프트 원문 (curate_with_gemini)")
P("상위 8개 토픽을 Gemini 2.5 Flash에 넘긴다. 입력은 한국어 프롬프트 + 각 토픽의 대표 제목/원문 요약/보도 매체 리스트다. "
  "출력은 엄격한 JSON 배열이며, 중간 텍스트는 금지한다.")

H2("프롬프트 본문 (요약 규칙)")
CODE(
    "오늘은 {today}입니다.\n"
    "아래는 AI 뉴스를 중복 보도 수 기준으로 정렬한 토픽 목록입니다.\n"
    "중복 보도 수는 이미 코드에서 계산된 정확한 값입니다.\n"
    "\n"
    "{articles_text}\n"
    "\n"
    "## 규칙 (엄격히 준수)\n"
    "1. AI 기술/제품/연구와 직접 관련 없는 토픽은 무조건 제외하세요. AI 회사가 언급되더라도\n"
    "   내용이 정치, 군사, 외교, 부동산, 데이터센터 물리적 위협 등이면 제외합니다.\n"
    "2. 남은 토픽 중 상위 {NEWS_COUNT}개를 선택하세요 (중복 보도 수가 많은 순).\n"
    "3. 각 토픽의 \"원문 요약\"에 있는 정보만 사용하세요. 원문에 없는 내용을 추가하거나\n"
    "   추측하지 마세요.\n"
    "4. 한국어로 번역/요약하되, 사실관계를 변경하지 마세요.\n"
    "5. \"topic_index\"는 위 목록의 [토픽 N]에서 N을 그대로 사용하세요.\n"
    "\n"
    "다른 텍스트 없이 JSON 배열만 출력하세요."
)

H2("JSON 스키마")
CODE(
    "{\n"
    '  "topic_index":   토픽 번호 (정수),\n'
    '  "tag_kr":        "LLM|이미지AI|로보틱스|규제|기업동향|연구 중 하나",\n'
    '  "tag_en":        "LLM|Image AI|Robotics|Regulation|Industry|Research 중 하나",\n'
    '  "tag_jp":        "LLM|画像AI|ロボティクス|規制|企業動向|研究 중 하나",\n'
    '  "title_kr":      "한국어 제목 (20자 이내)",\n'
    '  "title_en":      "영어 제목 (20자 이내)",\n'
    '  "title_jp":      "일본어 제목 (20자 이내)",\n'
    '  "summary_kr":    "한국어 핵심 요약 (2~3문장)",\n'
    '  "summary_en":    "영어 핵심 요약 (2~3문장)",\n'
    '  "summary_jp":    "일본어 핵심 요약 (2~3문장)",\n'
    '  "why_kr":        "왜 중요한가 (한국어 1문장)",\n'
    '  "why_en":        "왜 중요한가 (영어 1문장)",\n'
    '  "why_jp":        "왜 중요한가 (일본어 1문장)",\n'
    '  "company_domain": "기사 주체의 공식 도메인. 순수 도메인만 (예: openai.com,\n'
    '                    stanford.edu, softbank.jp). 모를 때만 빈 문자열"\n'
    "}"
)

H2("입력 포맷 (articles_text)")
CODE(
    "[토픽 N] (중복 보도: X개 매체)\n"
    "대표 제목: ...\n"
    "원문 요약: ...\n"
    "보도 매체: 매체1, 매체2, ..."
)

H2("재시도 정책")
BULLET("GEMINI_MAX_RETRIES = 3")
BULLET("지수 백오프: 10초 → 20초 → 40초 (코드 주석: '503 일시 과부하 대응')")
BULLET("JSON 파싱 실패, 검증 통과 건수 부족도 모두 재시도 대상")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. 환각 방지 장치
# ═══════════════════════════════════════════════════════════════
H1("8. 환각 방지 — 코드가 팩트를 덮어쓴다")
P("Gemini 응답을 그대로 쓰지 않는다. 팩트에 해당하는 필드는 Gemini가 JSON으로 내놓더라도 "
  "곧바로 코드가 원본 데이터로 덮어쓴다. 즉 Gemini는 '요약과 번역'만 책임지고, "
  "'링크/매체 수/원문 제목/발행일'은 코드가 책임진다.")

H2("덮어쓰는 필드 (정확 코드)")
CODE(
    'a["link"] = topic["link"]\n'
    'count = topic["source_count"]\n'
    'a["sources_kr"] = f"{count}개 매체 보도"\n'
    'a["sources_en"] = f"Covered by {count} sources"\n'
    'a["sources_jp"] = f"{count}社が報道"\n'
    'a["original_title"] = topic["title"]\n'
    'a["published"]     = topic.get("date")\n'
    'a["brand_key"]     = topic.get("brand_key", "")'
)

H2("왜 이렇게 하나")
BULLET("Gemini가 'URL을 그럴듯하게 조작'하거나 '매체 수를 부풀리는' 환각 사례가 실제로 있었음 (README Troubleshooting에도 언급)")
BULLET("매체 수는 클러스터링 결과로 이미 정확한 정수가 나와 있으니 LLM 판단이 필요 없다")
BULLET("원문 제목은 실제 RSS에서 받은 그대로가 '정답'이다")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. 파비콘 매칭
# ═══════════════════════════════════════════════════════════════
H1("9. 파비콘 매칭 — 어떤 회사 아이콘을 붙일지 결정")
P("파비콘은 카드 타이틀 왼쪽의 작은 회사 아이콘. 구현 철학은 "
  "'Gemini가 이미 알고 있는 도메인을 신뢰하되, 형식만 검증한다'이다. "
  "하드코딩 화이트리스트를 없애 스탠포드/소프트뱅크 등 사전 등록되지 않은 기관도 자동 커버된다.")

H2("매칭 우선순위")
p = doc.add_paragraph(style="List Number")
p.add_run("① Gemini company_domain 힌트").bold = True
p.add_run(" — URL/www/경로 자동 정리 후 정규식으로 형식 검증. 통과하면 즉시 사용.")
p = doc.add_paragraph(style="List Number")
p.add_run("② summary_kr 주어 매칭").bold = True
p.add_run(" — 한국어 요약의 가장 앞에 등장한 엔티티를 뽑는다. dict 순서가 아니라 실제 등장 위치(index) 기준.")
p = doc.add_paragraph(style="List Number")
p.add_run("③ title_kr 주어 매칭").bold = True
p.add_run(" — 같은 방식으로 한국어 제목 검사")
p = doc.add_paragraph(style="List Number")
p.add_run("④ 영문 원본 백업").bold = True
p.add_run(" — topic['title'] / topic['summary']에서 영문 substring 매칭")

H2("도메인 형식 검증 정규식")
CODE(r"^(?!-)[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?(?:\.[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?)+$")
P("통과 예: openai.com, stanford.edu, softbank.jp, group.softbank, x.ai")
P("탈락 예: 'Stanford University', 'www.stanford.edu/page', 'openai', '' (빈 문자열)")

H2("ENTITY_ALIASES 백업 사전")
P("Gemini가 도메인을 못 주는 경우를 위한 하드코딩 백업. 빅5부터 주요 스타트업/로봇/반도체까지 포함되며, "
  "한국어 별칭(오픈ai, 구글, 앤트로픽, 클로드, 마이크로소프트, 애플, 엔비디아, 아마존, 삼성, 퍼플렉시티, 허깅페이스 등)도 등록되어 있다. "
  "`메타`/`라마`/`엠에스`처럼 일반 단어와 충돌하는 것은 의도적으로 제외.")

H2("렌더링")
CODE('<img class="company-icon"\n'
     '     src="https://www.google.com/s2/favicons?domain={company}&sz=64"\n'
     '     alt="">')
P("Google Favicon API를 통해 해당 도메인의 실제 파비콘을 가져온다. 없는 도메인이면 기본 globe 아이콘이 표시된다.")

# ═══════════════════════════════════════════════════════════════
# 10. Tier 0 브랜드 컬러
# ═══════════════════════════════════════════════════════════════
H1("10. Tier 0 브랜드 컬러 — 공식 매체 카드 구분")
P("클러스터에 Tier 0 매체가 하나라도 포함되면 카드에 브랜드 색이 입혀진다. "
  "시각 신호는 세 겹이다: 상단 4px 액센트 바 + 배경 방사 글로우 + 테두리.")

H2("SOURCE_BRAND 매핑")
BULLET("OpenAI News → openai → 민트 그린 #10a37f")
BULLET("Google Research / Google DeepMind → google → Gemini 4색 그라데이션 (#4285f4, #ea4335, #fbbc04, #34a853)")
BULLET("Anthropic News → anthropic → Claude 주황 #cc785c")
BULLET("NVIDIA Blog → nvidia → 시그니처 그린 #76b900")
BULLET("Microsoft Research → microsoft → Azure 블루 #0078d4")

H2("브랜드 키 결정 로직")
P("한 클러스터에 여러 Tier 0 매체가 섞여 있으면 가중치가 가장 높은 매체의 brand_key를 사용. "
  "Tier 0가 하나도 없으면 brand_key는 빈 문자열이고 기본 카드 스타일이 적용된다.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. 다국어 + Edge TTS
# ═══════════════════════════════════════════════════════════════
H1("11. 다국어 지원과 Edge TTS")
P("카드 UI는 한국어/영어/일본어를 실시간 전환한다. 음성 브리핑도 3개 언어 모두 제공.")

H2("TTS 음성 모델")
BULLET("KR: ko-KR-SunHiNeural")
BULLET("EN: en-US-AriaNeural")
BULLET("JP: ja-JP-NanamiNeural")
P("Edge TTS로 기사별 음성 + 전체 연속 재생용 음성을 생성. 템포는 rate='+5%' (약간 빠르게).")

H2("재생 UX")
BULLET("각 카드 좌하단 ▶ 버튼으로 개별 재생")
BULLET("Spotify 스타일 하단 플레이어 바로 전체 순차 재생")
BULLET("언어 전환 시 재생 중이던 음성 자동 교체")

# ═══════════════════════════════════════════════════════════════
# 12. 배포와 자동화
# ═══════════════════════════════════════════════════════════════
H1("12. 배포와 자동화")

H2("실행 진입점")
BULLET("build.py — GitHub Actions에서 실행. signal_core.py의 파이프라인 전체를 호출하고, HTML/오디오를 Pages에 배포, archive에 커밋, 매체 통계 저장.")
BULLET("ai_briefing.py — 로컬 PC 실행용. Edge 앱 모드로 팝업 창에 브리핑을 띄운다. 매일 09:00 Windows 작업 스케줄러로 자동 실행.")

H2("GitHub Actions 스케줄")
BULLET("cron: 매일 UTC 20:00 (KST 05:00)")
BULLET("수동 실행: workflow_dispatch 지원")
BULLET("빌드 실패 시 build_error_html()로 에러 페이지 생성 후 배포")

H2("아카이브")
P("archive/ 디렉토리에 과거 브리핑 HTML을 날짜별로 저장. 각 빌드에서 git 커밋으로 올려 영구 보존한다. "
  "라이브 페이지 하단 'Past Briefings →' 링크에서 접근 가능.")

# ═══════════════════════════════════════════════════════════════
# 13. 매체 신뢰도 자동 학습
# ═══════════════════════════════════════════════════════════════
H1("13. 매체 신뢰도 자동 학습")
P("7일치 이상의 매체 통계가 쌓이면 SOURCE_WEIGHT를 자동으로 재분류한다. "
  "하드코딩된 가중치를 '실제 교차 보도율'로 덮어쓰는 것이 목적이다.")

H2("판정 공식")
CODE("rate = cross / total  # 기사의 몇 %가 타 매체에서도 보도됐는가\n"
     "if rate >= 0.25: Tier 1 (가중치 3)\n"
     "elif rate >= 0.10: Tier 2 (가중치 2)\n"
     "else: Tier 3 (가중치 1)")

H2("Tier 0은 자동 갱신에서 제외")
P("공식 1차 소스는 단독 보도가 본질이므로 교차 보도율이 낮아도 보호한다 "
  "(`if SOURCE_WEIGHT.get(src, _DEFAULT_WEIGHT) >= TIER_0_WEIGHT: continue`).")

H2("통계 기록")
BULLET("collect_source_metrics() — 매일 매체별 {total, cross, rate}를 계산")
BULLET("save_source_metrics() — metrics/YYYY-MM-DD.json으로 저장")
BULLET("update_source_weight() — 7일 이상 축적되면 자동 반영")

# ═══════════════════════════════════════════════════════════════
# 부록
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
H1("부록: 주요 상수 한눈에")
CODE(
    "NEWS_COUNT           = 5      # 최종 노출 뉴스 개수\n"
    "MODEL                = 'gemini-2.5-flash'\n"
    "EMBEDDING_MODEL      = 'gemini-embedding-001'\n"
    "SIMILARITY_THRESHOLD = 0.85   # 코사인 유사도 임계값 (0.80은 과묶음)\n"
    "EMBEDDING_BATCH_SIZE = 100    # 배치당 텍스트 수\n"
    "EMBEDDING_BATCH_SLEEP= 20     # 배치 간 대기(초) — RPM/TPM 여유\n"
    "GEMINI_MAX_RETRIES   = 3      # 실패 시 재시도 횟수\n"
    "TIER_0_WEIGHT        = 4      # Tier 0 = 공식 1차 소스\n"
    "RSS cutoff (기본)    = 24h    # 부족 시 36h로 확장\n"
    "TTS_VOICES           = {\n"
    "    'kr': 'ko-KR-SunHiNeural',\n"
    "    'en': 'en-US-AriaNeural',\n"
    "    'jp': 'ja-JP-NanamiNeural',\n"
    "}"
)

H1("부록: 전체 파일 구조")
BULLET("signal_core.py — 공통 코어 (RSS, 임베딩, Gemini, HTML, TTS) — 1,326줄")
BULLET("build.py — GitHub Actions 빌드 래퍼 (Pages 배포 + 아카이브 + 통계)")
BULLET("ai_briefing.py — 로컬 실행 래퍼 (Edge 앱 창 팝업)")
BULLET("archive/ — 과거 브리핑 HTML (git 커밋으로 영구 보존)")
BULLET("metrics/ — 매체별 교차 보도율 통계 (날짜별 JSON)")
BULLET("favicon-{32,180,512}.png — 사이트 파비콘")

# 저장
out_path = r"C:\Users\user\ai_briefing\THE_SIGNAL_상세설명서.docx"
doc.save(out_path)
print(f"saved: {out_path}")
